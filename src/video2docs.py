#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Video to Document Converter

This script converts YouTube videos or local video files to document formats (ODT, DOCX, PDF).
It extracts text from speech, identifies slides/images, and uses LLMs to organize the content.
"""

import os
import argparse
import tempfile
from pathlib import Path
from typing import List, Dict, Tuple, Optional, Union
import logging
import re
import time
from dotenv import load_dotenv

# Video processing
import cv2
from pytubefix import YouTube
from pytubefix.cli import on_progress
from moviepy import VideoFileClip

# Audio processing
import speech_recognition as sr
from pydub import AudioSegment

# Image processing
from PIL import Image
import numpy as np
from skimage.metrics import structural_similarity as ssim

# Document generation
import docx
from docx.shared import Inches
from odf.opendocument import OpenDocumentText
from odf.text import P
from odf.style import Style, TextProperties
from odf.draw import Frame, Image as ODFImage
from fpdf import FPDF

# LLM and AI
import torch
from transformers import pipeline, AutoModelForSeq2SeqLM, AutoTokenizer
from langchain_community.llms import OpenAI
from langchain_community.llms.huggingface_pipeline import HuggingFacePipeline
from langchain.prompts import PromptTemplate
from langchain.schema.runnable import RunnablePassthrough

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()


class VideoProcessor:
    """Handles video downloading and processing."""

    def __init__(self, temp_dir: str = None):
        """Initialize the video processor.

        Args:
            temp_dir: Directory to store temporary files
        """
        self.temp_dir = temp_dir or tempfile.mkdtemp()
        logger.info(f"Using temporary directory: {self.temp_dir}")

    def download_youtube_video(self, url: str, max_retries: int = 3) -> str:
        """Download a YouTube video.

        Args:
            url: YouTube video URL
            max_retries: Maximum number of retry attempts

        Returns:
            Path to the downloaded video file
        """
        logger.info(f"Downloading YouTube video: {url}")

        # List of common YouTube domains to try if the original URL fails
        youtube_domains = [
            "www.youtube.com",
            "youtube.com",
            "youtu.be",
            "m.youtube.com",
        ]

        # Extract video ID from URL
        video_id = None
        if "youtube.com/watch" in url:
            # Format: https://www.youtube.com/watch?v=VIDEO_ID
            query_params = url.split("?")[1].split("&")
            for param in query_params:
                if param.startswith("v="):
                    video_id = param[2:]
                    break
        elif "youtu.be/" in url:
            # Format: https://youtu.be/VIDEO_ID
            video_id = url.split("youtu.be/")[1].split("?")[0]

        if not video_id:
            error_msg = f"Could not extract video ID from URL: {url}"
            logger.error(error_msg)
            raise ValueError(error_msg)

        # Try downloading with the original URL first
        urls_to_try = [url]

        # Add alternative URLs with different domains
        for domain in youtube_domains:
            if domain not in url:
                alt_url = f"https://{domain}/watch?v={video_id}"
                if alt_url != url:
                    urls_to_try.append(alt_url)

        last_exception = None
        for retry_count in range(max_retries):
            for try_url in urls_to_try:
                try:
                    logger.info(f"Attempt {retry_count + 1}/{max_retries} with URL: {try_url}")
                    yt = YouTube(try_url, on_progress_callback=on_progress)
                    logger.info(yt.title)
                    video = yt.streams.get_highest_resolution()

                    if not video:
                        logger.warning(f"No suitable video streams found for {try_url}")
                        continue

                    output_path = video.download(output_path=self.temp_dir)
                    logger.info(f"Downloaded video to: {output_path}")
                    return output_path
                except Exception as e:
                    logger.warning(f"Error downloading YouTube video from {try_url}: {e}")
                    last_exception = e
                    # Continue to the next URL or retry

            # Wait before retrying
            if retry_count < max_retries - 1:
                wait_time = 2 ** retry_count  # Exponential backoff
                logger.info(f"Waiting {wait_time} seconds before next retry...")
                time.sleep(wait_time)

        # If we get here, all attempts failed
        error_msg = f"Failed to download YouTube video after {max_retries} attempts with multiple URLs. Last error: {last_exception}"
        logger.error(error_msg)
        raise Exception(error_msg)

    def extract_audio(self, video_path: str) -> str:
        """Extract audio from video file.

        Args:
            video_path: Path to the video file

        Returns:
            Path to the extracted audio file
        """
        logger.info(f"Extracting audio from: {video_path}")
        try:
            video = VideoFileClip(video_path)
            audio_path = os.path.join(self.temp_dir, "audio.wav")
            video.audio.write_audiofile(audio_path, codec='pcm_s16le')
            logger.info(f"Extracted audio to: {audio_path}")
            return audio_path
        except Exception as e:
            logger.error(f"Error extracting audio: {e}")
            raise

    def extract_frames(self, video_path: str, interval: float = 1.0) -> List[Tuple[float, str]]:
        """Extract frames from video at specified intervals.

        Args:
            video_path: Path to the video file
            interval: Interval between frames in seconds

        Returns:
            List of tuples containing (timestamp, frame_path)
        """
        logger.info(f"Extracting frames from: {video_path}")
        frames = []
        try:
            video = cv2.VideoCapture(video_path)
            fps = video.get(cv2.CAP_PROP_FPS)
            frame_interval = int(fps * interval)

            success, frame = video.read()
            count = 0
            frame_count = 0

            while success:
                if count % frame_interval == 0:
                    timestamp = count / fps
                    frame_path = os.path.join(self.temp_dir, f"frame_{frame_count:04d}.jpg")
                    cv2.imwrite(frame_path, frame)
                    frames.append((timestamp, frame_path))
                    frame_count += 1

                success, frame = video.read()
                count += 1

            video.release()
            logger.info(f"Extracted {len(frames)} frames")
            return frames
        except Exception as e:
            logger.error(f"Error extracting frames: {e}")
            raise

    def detect_slides(self, frames: List[Tuple[float, str]], threshold: float = 0.8) -> List[Tuple[float, str]]:
        """Detect slides/images in the extracted frames.

        Args:
            frames: List of (timestamp, frame_path) tuples
            threshold: Similarity threshold for slide detection

        Returns:
            List of (timestamp, frame_path) tuples for detected slides
        """
        logger.info("Detecting slides in frames")
        if not frames:
            return []

        slides = [frames[0]]  # First frame is always a slide

        for i in range(1, len(frames)):
            prev_img = cv2.imread(slides[-1][1], cv2.IMREAD_GRAYSCALE)
            curr_img = cv2.imread(frames[i][1], cv2.IMREAD_GRAYSCALE)

            # Resize images to same dimensions if needed
            if prev_img.shape != curr_img.shape:
                curr_img = cv2.resize(curr_img, (prev_img.shape[1], prev_img.shape[0]))

            # Calculate similarity
            similarity = ssim(prev_img, curr_img)

            if similarity < threshold:
                slides.append(frames[i])
                logger.debug(f"Detected new slide at {frames[i][0]:.2f}s (similarity: {similarity:.2f})")

        logger.info(f"Detected {len(slides)} slides")
        return slides


class AudioProcessor:
    """Handles audio transcription."""

    def __init__(self, use_gpu: bool = True):
        """Initialize the audio processor.

        Args:
            use_gpu: Whether to use GPU for processing
        """
        self.use_gpu = use_gpu and torch.cuda.is_available()
        self.recognizer = sr.Recognizer()
        logger.info(f"Initialized audio processor (GPU: {self.use_gpu})")

    def transcribe_audio(self, audio_path: str, chunk_size: int = 60000) -> List[Dict[str, Union[str, float, float]]]:
        """Transcribe audio file to text.

        Args:
            audio_path: Path to the audio file
            chunk_size: Size of audio chunks in milliseconds

        Returns:
            List of dictionaries with text, start_time, and end_time
        """
        logger.info(f"Transcribing audio: {audio_path}")
        try:
            audio = AudioSegment.from_file(audio_path)
            duration = len(audio)
            chunks = []

            for start_time in range(0, duration, chunk_size):
                end_time = min(start_time + chunk_size, duration)
                chunk = audio[start_time:end_time]

                # Save chunk to temporary file
                chunk_path = f"{audio_path}_chunk_{start_time}_{end_time}.wav"
                try:
                    chunk.export(chunk_path, format="wav")

                    # Transcribe chunk
                    with sr.AudioFile(chunk_path) as source:
                        audio_data = self.recognizer.record(source)
                        logger.info(f"Transcribing chunk from {start_time} to {end_time} ms")
                        try:
                            text = self.recognizer.recognize_google(audio_data)
                            logger.info(f"Transcribed chunk: {text}")
                        except sr.UnknownValueError:
                            logger.warning(f"Speech not understood for chunk {start_time}-{end_time} ms; leaving text empty.")
                            text = ""
                        except sr.RequestError as re_err:
                            logger.error(f"Speech recognition service error for chunk {start_time}-{end_time} ms: {re_err}")
                            text = ""

                    chunks.append({
                        "text": text,
                        "start_time": start_time / 1000.0,  # Convert to seconds
                        "end_time": end_time / 1000.0  # Convert to seconds
                    })
                finally:
                    # Clean up temporary file
                    if os.path.exists(chunk_path):
                        try:
                            os.remove(chunk_path)
                        except Exception as cleanup_err:
                            logger.warning(f"Failed to remove temp chunk file {chunk_path}: {cleanup_err}")

            logger.info(f"Transcribed {len(chunks)} audio chunks")
            return chunks
        except Exception as e:
            logger.error(f"Error transcribing audio: {e}")
            raise


class LLMProcessor:
    """Handles LLM processing for content organization."""

    def __init__(self, model_name: str = None, use_openai: bool = False, use_gpu: bool = True):
        """Initialize the LLM processor.

        Args:
            model_name: Name of the HuggingFace model to use
            use_openai: Whether to use OpenAI API
            use_gpu: Whether to use GPU for processing
        """
        self.use_openai = use_openai
        self.use_gpu = use_gpu and torch.cuda.is_available()

        if use_openai and os.getenv("OPENAI_API_KEY"):
            logger.info("Using OpenAI for LLM processing")
            self.llm = OpenAI(temperature=0.1)
        else:
            # Default to a HuggingFace model
            self.model_name = model_name or "google/flan-t5-large"
            logger.info(f"Using HuggingFace model: {self.model_name}")

            # Set device based on GPU availability and preference
            device = "cuda" if self.use_gpu else "cpu"
            logger.info(f"Using device: {device}")

            # Download and load the model locally
            logger.info(f"Downloading and loading model: {self.model_name}")
            try:
                model = AutoModelForSeq2SeqLM.from_pretrained(self.model_name)
                tokenizer = AutoTokenizer.from_pretrained(self.model_name)

                # Move model to GPU if available
                model = model.to(device)
                logger.info(f"Successfully loaded model {self.model_name} to {device}")
            except Exception as e:
                logger.error(f"Error loading model {self.model_name}: {e}")
                logger.error("Please check your internet connection and ensure you have enough disk space.")
                logger.error("If the error persists, try using a different model or check if the model is available on Hugging Face Hub.")
                raise RuntimeError(f"Failed to load model {self.model_name}: {e}") from e

            # Create a text generation pipeline
            text_generation_pipeline = pipeline(
                "text2text-generation",
                model=model,
                tokenizer=tokenizer,
                device=0 if device == "cuda" else -1,
                max_length=512,
                temperature=0.1,
                num_return_sequences=1,  # Only return one sequence
                do_sample=True  # Enable sampling for more creative outputs
            )

            # Create a LangChain pipeline
            self.llm = HuggingFacePipeline(pipeline=text_generation_pipeline)

    def organize_content(self, transcription: List[Dict], slides: List[Tuple[float, str]]) -> Dict:
        """Organize content using LLM.

        Args:
            transcription: List of transcription chunks
            slides: List of detected slides

        Returns:
            Organized document structure
        """
        logger.info("Organizing content with LLM")

        # Combine transcription into a single text
        full_text = " ".join([chunk["text"] for chunk in transcription])

        # Create prompt for document organization
        prompt = PromptTemplate(
            input_variables=["text", "num_slides"],
            template="""
            You are an AI assistant that organizes video content into a structured document.

            The video contains the following transcribed text:
            {text}

            The video also contains {num_slides} slides or important images.

            Please organize this content into a well-structured document with:
            1. A title
            2. An executive summary
            3. Main sections with headings
            4. Bullet points for key information
            5. Indications where slides should be inserted (marked as [SLIDE X])

            Format your response as a JSON with the following structure:
            {{
                "title": "Document Title",
                "summary": "Executive summary text",
                "sections": [
                    {{
                        "heading": "Section Heading",
                        "content": "Section content with [SLIDE X] markers where appropriate",
                        "bullet_points": ["Point 1", "Point 2"]
                    }}
                ]
            }}
            """
        )

        # Create formatted prompt - limit text to fit within model's token limit (512 tokens)
        # Assuming ~4 chars per token, limit to ~1500 chars to leave room for prompt template
        formatted_prompt = prompt.format(text=full_text[:1500], num_slides=len(slides))

        # Invoke the LLM
        raw_result = self.llm.invoke(formatted_prompt)

        # Extract text from the result (HuggingFacePipeline returns a list of dictionaries)
        if isinstance(raw_result, list) and len(raw_result) > 0 and 'generated_text' in raw_result[0]:
            result = raw_result[0]['generated_text']
        else:
            result = raw_result  # Fallback to original format for compatibility

        logger.debug(f"Raw LLM result type: {type(raw_result)}")
        logger.debug(f"Processed result: {result[:100]}...")

        # Parse the result
        try:
            # Extract JSON from the response (it might be surrounded by markdown code blocks)
            json_match = re.search(r'```json\n(.*?)\n```', result, re.DOTALL)
            if json_match:
                result = json_match.group(1)

            import json
            document_structure = json.loads(result)
            logger.info("Successfully organized content")
            return document_structure
        except Exception as e:
            logger.error(f"Error parsing LLM response: {e}")
            logger.debug(f"Raw LLM response: {result}")

            # Return a basic structure if parsing fails
            return {
                "title": "Transcribed Video",
                "summary": full_text[:500] + "...",
                "sections": [
                    {
                        "heading": "Full Transcription",
                        "content": full_text,
                        "bullet_points": []
                    }
                ]
            }


class DocumentGenerator:
    """Generates documents in various formats."""

    def __init__(self, output_dir: str = "."):
        """Initialize the document generator.

        Args:
            output_dir: Directory to save output documents
        """
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        logger.info(f"Document generator initialized (output dir: {output_dir})")

    def generate_docx(self, content: Dict, slides: List[Tuple[float, str]], output_path: str) -> str:
        """Generate a DOCX document.

        Args:
            content: Organized document content
            slides: List of detected slides
            output_path: Path to save the document

        Returns:
            Path to the generated document
        """
        logger.info(f"Generating DOCX document: {output_path}")
        doc = docx.Document()

        # Add title
        doc.add_heading(content["title"], level=0)

        # Add summary
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(content["summary"])

        # Add a section for slides/images if available
        if slides:
            doc.add_heading("Slides/Images", level=1)
            for i, (timestamp, slide_path) in enumerate(slides):
                try:
                    # Add a caption for the slide
                    doc.add_paragraph(f"Slide {i+1} (Timestamp: {timestamp:.2f}s)")
                    # Add the image
                    doc.add_picture(slide_path, width=Inches(6))
                    # Add some space after each image
                    doc.add_paragraph()
                except Exception as e:
                    logger.error(f"Error adding slide {i} from {slide_path}: {e}")

        # Add sections
        for section in content["sections"]:
            doc.add_heading(section["heading"], level=1)

            # Process content with slide markers
            content_parts = re.split(r'(\[SLIDE \d+\])', section["content"])
            for part in content_parts:
                slide_match = re.match(r'\[SLIDE (\d+)\]', part)
                if slide_match:
                    slide_index = int(slide_match.group(1))
                    if 0 <= slide_index < len(slides):
                        try:
                            doc.add_picture(slides[slide_index][1], width=Inches(6))
                        except Exception as e:
                            logger.error(f"Error adding slide {slide_index} from {slides[slide_index][1]}: {e}")
                else:
                    if part.strip():
                        doc.add_paragraph(part)

            # Add bullet points
            if section.get("bullet_points"):
                for point in section["bullet_points"]:
                    doc.add_paragraph(point, style='List Bullet')

        # Save document
        doc.save(output_path)
        logger.info(f"DOCX document saved to: {output_path}")
        return output_path

    def generate_odt(self, content: Dict, slides: List[Tuple[float, str]], output_path: str) -> str:
        """Generate an ODT document.

        Args:
            content: Organized document content
            slides: List of detected slides
            output_path: Path to save the document

        Returns:
            Path to the generated document
        """
        logger.info(f"Generating ODT document: {output_path}")
        doc = OpenDocumentText()

        # Add styles
        heading_style = Style(name="Heading", family="paragraph")
        heading_style.addElement(TextProperties(attributes={'fontsize': "16pt", 'fontweight': "bold"}))
        doc.styles.addElement(heading_style)

        # Add title
        title = P(stylename=heading_style)
        title.addText(content["title"])
        doc.text.addElement(title)

        # Add summary
        summary_heading = P(stylename=heading_style)
        summary_heading.addText("Executive Summary")
        doc.text.addElement(summary_heading)

        summary = P()
        summary.addText(content["summary"])
        doc.text.addElement(summary)

        # Add a section for slides/images if available
        if slides:
            slides_heading = P(stylename=heading_style)
            slides_heading.addText("Slides/Images")
            doc.text.addElement(slides_heading)

            for i, (timestamp, slide_path) in enumerate(slides):
                try:
                    # Add a caption for the slide
                    caption = P()
                    caption.addText(f"Slide {i+1} (Timestamp: {timestamp:.2f}s)")
                    doc.text.addElement(caption)

                    # Add image
                    frame = Frame(width="6in", height="4in")
                    img = ODFImage(href=slide_path)
                    frame.addElement(img)
                    doc.text.addElement(frame)

                    # Add some space after each image
                    spacer = P()
                    doc.text.addElement(spacer)
                except Exception as e:
                    logger.error(f"Error adding slide {i} from {slide_path}: {e}")

        # Add sections
        for section in content["sections"]:
            section_heading = P(stylename=heading_style)
            section_heading.addText(section["heading"])
            doc.text.addElement(section_heading)

            # Process content with slide markers
            content_parts = re.split(r'(\[SLIDE \d+\])', section["content"])
            for part in content_parts:
                slide_match = re.match(r'\[SLIDE (\d+)\]', part)
                if slide_match:
                    slide_index = int(slide_match.group(1))
                    if 0 <= slide_index < len(slides):
                        try:
                            # Add image
                            frame = Frame(width="6in", height="4in")
                            img = ODFImage(href=slides[slide_index][1])
                            frame.addElement(img)
                            doc.text.addElement(frame)
                        except Exception as e:
                            logger.error(f"Error adding slide {slide_index} from {slides[slide_index][1]}: {e}")
                else:
                    if part.strip():
                        p = P()
                        p.addText(part)
                        doc.text.addElement(p)

            # Add bullet points
            if section.get("bullet_points"):
                for point in section["bullet_points"]:
                    bullet = P(stylename="List")
                    bullet.addText("• " + point)
                    doc.text.addElement(bullet)

        # Save document
        doc.save(output_path)
        logger.info(f"ODT document saved to: {output_path}")
        return output_path

    def generate_pdf(self, content: Dict, slides: List[Tuple[float, str]], output_path: str) -> str:
        """Generate a PDF document.

        Args:
            content: Organized document content
            slides: List of detected slides
            output_path: Path to save the document

        Returns:
            Path to the generated document
        """
        logger.info(f"Generating PDF document: {output_path}")
        pdf = FPDF()
        pdf.add_page()

        # Set up fonts
        pdf.set_font("Arial", "B", 16)

        # Add title
        pdf.cell(0, 10, content["title"], ln=True, align="C")
        pdf.ln(10)

        # Add summary
        pdf.set_font("Arial", "B", 14)
        pdf.cell(0, 10, "Executive Summary", ln=True)
        pdf.set_font("Arial", "", 12)
        pdf.multi_cell(0, 10, content["summary"])
        pdf.ln(10)

        # Add a section for slides/images if available
        if slides:
            pdf.set_font("Arial", "B", 14)
            pdf.cell(0, 10, "Slides/Images", ln=True)
            pdf.set_font("Arial", "", 12)

            for i, (timestamp, slide_path) in enumerate(slides):
                try:
                    # Add a caption for the slide
                    pdf.multi_cell(0, 10, f"Slide {i+1} (Timestamp: {timestamp:.2f}s)")

                    # Add image
                    pdf.image(slide_path, x=10, w=190)

                    # Add some space after each image
                    pdf.ln(5)
                except Exception as e:
                    logger.error(f"Error adding slide {i} from {slide_path}: {e}")

            pdf.ln(10)

        # Add sections
        for section in content["sections"]:
            pdf.set_font("Arial", "B", 14)
            pdf.cell(0, 10, section["heading"], ln=True)
            pdf.set_font("Arial", "", 12)

            # Process content with slide markers
            content_parts = re.split(r'(\[SLIDE \d+\])', section["content"])
            for part in content_parts:
                slide_match = re.match(r'\[SLIDE (\d+)\]', part)
                if slide_match:
                    slide_index = int(slide_match.group(1))
                    if 0 <= slide_index < len(slides):
                        try:
                            # Add image
                            pdf.image(slides[slide_index][1], x=10, w=190)
                        except Exception as e:
                            logger.error(f"Error adding slide {slide_index} from {slides[slide_index][1]}: {e}")
                else:
                    if part.strip():
                        pdf.multi_cell(0, 10, part)

            # Add bullet points
            if section.get("bullet_points"):
                for point in section["bullet_points"]:
                    pdf.cell(10, 10, "•", ln=0)
                    pdf.multi_cell(0, 10, point)

        # Save document
        pdf.output(output_path)
        logger.info(f"PDF document saved to: {output_path}")
        return output_path


class Video2Docs:
    """Main class for video to document conversion."""

    def __init__(self, output_dir: str = "output", temp_dir: str = None, use_gpu: bool = True):
        """Initialize the converter.

        Args:
            output_dir: Directory to save output documents
            temp_dir: Directory for temporary files
            use_gpu: Whether to use GPU for processing
        """
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

        self.temp_dir = temp_dir or os.path.join(output_dir, "temp")
        os.makedirs(self.temp_dir, exist_ok=True)

        self.use_gpu = use_gpu and torch.cuda.is_available()
        if self.use_gpu:
            logger.info("GPU is available and will be used")
        else:
            logger.info("GPU is not available, using CPU")

        # Initialize components
        self.video_processor = VideoProcessor(temp_dir=self.temp_dir)
        self.audio_processor = AudioProcessor(use_gpu=self.use_gpu)
        self.llm_processor = LLMProcessor(use_openai=False, use_gpu=self.use_gpu)
        self.document_generator = DocumentGenerator(output_dir=self.output_dir)

    def process(self, input_path: str, output_format: str = "docx", output_name: Optional[str] = None, progress_callback=None, cancel_event=None) -> str:
        """Process a video and convert it to a document.

        Args:
            input_path: Path or URL to the video
            output_format: Output document format (docx, odt, pdf)
            output_name: Optional base name (without extension) for the output file. If not provided,
                the name is derived from the input video file or video title.

        Returns:
            Path to the generated document
        """
        start_time = time.time()
        logger.info(f"Starting conversion of: {input_path}")

        # Progress/ETA helpers
        weights = {
            "download": 0.10,
            "extract_audio": 0.05,
            "extract_frames": 0.25,
            "detect_slides": 0.10,
            "transcribe": 0.30,
            "organize_content": 0.05,
            "generate_document": 0.15,
        }
        base = 0.0

        def report(step: str, step_progress: float):
            if progress_callback:
                total_progress = (base + weights.get(step, 0.0) * float(step_progress or 0.0)) * 100.0
                elapsed = time.time() - start_time
                eta = int(elapsed * (100.0 - total_progress) / max(total_progress, 1e-6)) if total_progress > 0 else None
                try:
                    progress_callback(step, round(total_progress, 1), eta)
                except Exception:
                    pass

        def check_cancel():
            if cancel_event is not None:
                try:
                    if cancel_event.is_set():
                        from .jobs import CancelledError
                        raise CancelledError()
                except AttributeError:
                    # Non-standard event
                    pass

        try:
            # Step 1: Get the video file
            report("download", 0.0)
            if input_path.startswith(("http://", "https://")) and "youtube" in input_path:
                try:
                    video_path = self.video_processor.download_youtube_video(input_path)
                    base_name = os.path.basename(video_path).split(".")[0]
                except Exception as youtube_error:
                    # Provide a more helpful error message for YouTube download failures
                    error_msg = str(youtube_error)
                    if "HTTP Error 400: Bad Request" in error_msg:
                        raise ValueError(
                            f"Failed to download YouTube video: {input_path}\n"
                            f"This could be due to one of the following reasons:\n"
                            f"1. The video might be unavailable, private, or age-restricted\n"
                            f"2. YouTube may have changed its API, affecting the pytube library\n"
                            f"3. There might be network restrictions or proxy issues\n\n"
                            f"Possible solutions:\n"
                            f"- Try downloading a local video file instead of a YouTube URL\n"
                            f"- Check if pytube has been updated to address YouTube API changes\n"
                            f"- Try running the script on a different network\n"
                            f"- Consider using youtube-dl or yt-dlp as alternative download methods\n\n"
                            f"Original error: {error_msg}"
                        ) from youtube_error
                    else:
                        # For other YouTube-related errors
                        raise ValueError(
                            f"Error downloading YouTube video: {input_path}\n"
                            f"Please check your internet connection and try again.\n"
                            f"Original error: {error_msg}"
                        ) from youtube_error
            else:
                # Local video file
                if not os.path.exists(input_path):
                    raise FileNotFoundError(f"Video file not found: {input_path}")

                video_path = input_path
                base_name = os.path.basename(video_path).split(".")[0]

            # Download complete
            base += weights.get("download", 0.0)
            report("download", 1.0)
            check_cancel()

            # Override base name if a custom output name is provided
            if output_name:
                base_name = os.path.splitext(output_name)[0]

            # Step 2: Extract audio
            report("extract_audio", 0.0)
            audio_path = self.video_processor.extract_audio(video_path)
            base += weights.get("extract_audio", 0.0)
            report("extract_audio", 1.0)
            check_cancel()

            # Step 3: Extract frames and detect slides
            report("extract_frames", 0.0)
            frames = self.video_processor.extract_frames(video_path)
            base += weights.get("extract_frames", 0.0)
            report("extract_frames", 1.0)
            check_cancel()

            report("detect_slides", 0.0)
            slides = self.video_processor.detect_slides(frames)
            base += weights.get("detect_slides", 0.0)
            report("detect_slides", 1.0)
            check_cancel()

            # Step 4: Transcribe audio
            report("transcribe", 0.0)
            transcription = self.audio_processor.transcribe_audio(audio_path)
            base += weights.get("transcribe", 0.0)
            report("transcribe", 1.0)
            check_cancel()

            # Step 5: Organize content with LLM
            report("organize_content", 0.0)
            content = self.llm_processor.organize_content(transcription, slides)
            base += weights.get("organize_content", 0.0)
            report("organize_content", 1.0)
            check_cancel()

            # Step 6: Generate document
            output_path = os.path.join(self.output_dir, f"{base_name}.{output_format}")

            report("generate_document", 0.0)
            if output_format.lower() == "docx":
                result_path = self.document_generator.generate_docx(content, slides, output_path)
            elif output_format.lower() == "odt":
                result_path = self.document_generator.generate_odt(content, slides, output_path)
            elif output_format.lower() == "pdf":
                result_path = self.document_generator.generate_pdf(content, slides, output_path)
            else:
                raise ValueError(f"Unsupported output format: {output_format}")

            base += weights.get("generate_document", 0.0)
            report("generate_document", 1.0)
            check_cancel()

            elapsed_time = time.time() - start_time
            logger.info(f"Conversion completed in {elapsed_time:.2f} seconds")
            logger.info(f"Output document: {result_path}")

            return result_path

        except ValueError as ve:
            # Re-raise ValueError exceptions (including our custom YouTube errors)
            logger.error(f"Error during conversion: {ve}")
            raise
        except FileNotFoundError as fnf:
            # Re-raise FileNotFoundError exceptions
            logger.error(f"Error during conversion: {fnf}")
            raise
        except Exception as e:
            # Generic error handling for other exceptions
            logger.error(f"Error during conversion: {e}")
            raise


def main():
    """Main entry point."""
    parser = argparse.ArgumentParser(description="Convert video to document")
    parser.add_argument("input", help="YouTube URL or path to local video file")
    parser.add_argument(
        "--format", "-f", 
        choices=["docx", "odt", "pdf"], 
        default="docx",
        help="Output document format"
    )
    parser.add_argument(
        "--output-dir", "-o",
        default="output",
        help="Directory to save output documents"
    )
    parser.add_argument(
        "--temp-dir", "-t",
        default=None,
        help="Directory for temporary files"
    )
    parser.add_argument(
        "--no-gpu",
        action="store_true",
        help="Disable GPU usage"
    )
    parser.add_argument(
        "--verbose", "-v",
        action="store_true",
        help="Enable verbose logging"
    )

    args = parser.parse_args()

    # Configure logging
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)

    # Create converter
    converter = Video2Docs(
        output_dir=args.output_dir,
        temp_dir=args.temp_dir,
        use_gpu=not args.no_gpu
    )

    # Process video
    try:
        output_path = converter.process(args.input, args.format)
        print(f"Document generated: {output_path}")
        return 0
    except Exception as e:
        logger.error(f"Error: {e}")
        return 1


if __name__ == "__main__":
    exit(main())
